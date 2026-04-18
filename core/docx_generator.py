from __future__ import annotations

import io
import random
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


TYPE_LABELS = {
    "subjective_short": "주관식 단답형",
    "subjective_long": "주관식 서술형",
    "multiple_choice_single": "객관식 단일선택",
    "multiple_choice_multi": "객관식 다중선택",
}


def _question_score(question: dict) -> int:
    try:
        return int(question.get("points", 5) or 5)
    except (TypeError, ValueError):
        return 5


def _normalize_answer(value):
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value or "").strip()
    return text


def _set_base_style(document: Document):
    style = document.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)


def _add_header(document: Document, settings: dict, version_label: str | None = None):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(str(settings.get("course_name") or settings.get("institution_name") or "Teach-On 시험지"))
    run.bold = True
    run.font.size = Pt(16)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = []
    if settings.get("institution_name"):
        bits.append(f"기관: {settings['institution_name']}")
    if settings.get("exam_date"):
        bits.append(f"시험일: {settings['exam_date']}")
    if settings.get("time_limit_minutes"):
        bits.append(f"제한시간: {settings['time_limit_minutes']}분")
    if version_label:
        bits.append(f"버전: {version_label}")
    meta.add_run(" | ".join(bits) or "자동 생성 시험지")


def _add_question_block(document: Document, index: int, question: dict, include_answers: bool):
    heading = document.add_paragraph()
    run = heading.add_run(f"{index}. {question.get('prompt', '').strip()}")
    run.bold = True
    heading.paragraph_format.space_after = Pt(4)

    sub = document.add_paragraph()
    sub.add_run(
        f"[{TYPE_LABELS.get(question.get('type'), question.get('type', '문항'))}] "
        f"난이도 {question.get('difficulty', '중')} / {_question_score(question)}점 / "
        f"출처 {question.get('source_pages', '-')}"
    )
    sub.paragraph_format.space_after = Pt(4)

    choices = question.get("choices") or []
    for idx, choice in enumerate(choices, start=1):
        line = document.add_paragraph(style="List Bullet")
        line.add_run(f"{idx}) {choice}")

    if include_answers:
        answer = _normalize_answer(question.get("answer"))
        answer_text = ", ".join(answer) if isinstance(answer, list) else answer
        p = document.add_paragraph()
        p.add_run(f"정답: {answer_text}").bold = True
        if question.get("explanation"):
            e = document.add_paragraph()
            e.add_run(f"해설: {question.get('explanation')}")


def _build_docx(questions: list[dict], settings: dict, include_answers: bool, version_label: str | None = None) -> io.BytesIO:
    document = Document()
    _set_base_style(document)
    _add_header(document, settings, version_label=version_label)
    document.add_paragraph("")
    for index, question in enumerate(questions, start=1):
        _add_question_block(document, index, question, include_answers=include_answers)
        if index != len(questions):
            document.add_paragraph("")
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return buf


def build_exam_artifacts(questions: list[dict], settings: dict) -> dict[str, io.BytesIO]:
    normalized = [deepcopy(question) for question in questions or []]
    artifacts = {
        "exam": _build_docx(normalized, settings, include_answers=False),
        "answer": _build_docx(normalized, settings, include_answers=True),
    }
    if settings.get("shuffle_versions"):
        seed_base = str(settings.get("shuffle_seed") or settings.get("uid") or "teachon")
        qa = deepcopy(normalized)
        qb = deepcopy(normalized)
        random.Random(seed_base + "A").shuffle(qa)
        random.Random(seed_base + "B").shuffle(qb)
        artifacts["exam_a"] = _build_docx(qa, settings, include_answers=False, version_label="A형")
        artifacts["exam_b"] = _build_docx(qb, settings, include_answers=False, version_label="B형")
    return artifacts
